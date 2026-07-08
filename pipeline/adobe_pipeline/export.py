"""Exporters: CSV (the scalable dataset format), JSON, and a clean, presentable Word doc."""
import csv
import json


COLUMNS = ["kind", "source", "external_id", "title", "company", "vertical", "adobe_tools",
           "budget", "salary", "location", "posted_at", "url", "description"]


def _rows(conn, kind=None):
    sql = "SELECT * FROM items"
    args = ()
    if kind:
        sql += " WHERE kind=?"
        args = (kind,)
    sql += " ORDER BY vertical, length(description) DESC"
    return conn.execute(sql, args).fetchall()


def export_csv(conn, path, kind=None):
    rows = _rows(conn, kind)
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(COLUMNS)
        for r in rows:
            tools = "; ".join(json.loads(r["adobe_tools"] or "[]"))
            w.writerow([r["kind"], r["source"], r["external_id"], r["title"], r["company"],
                        r["vertical"], tools, r["budget"], r["salary"], r["location"],
                        r["posted_at"], r["url"], r["description"]])
    return len(rows)


def export_json(conn, path, kind=None):
    rows = _rows(conn, kind)
    out = []
    for r in rows:
        d = dict(r)
        d["adobe_tools"] = json.loads(r["adobe_tools"] or "[]")
        d.pop("raw_json", None)
        out.append(d)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    return len(out)


# ---------------- Word export ----------------
def _add_hyperlink(paragraph, url, text, color="1155CC"):
    from docx.oxml.shared import OxmlElement, qn
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    paragraph._p.append(link)


def export_docx(conn, path, kind="task", limit=None):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    INK = RGBColor(0x16, 0x24, 0x3F); ACCENT = RGBColor(0xC8, 0x47, 0x2B)
    GREY = RGBColor(0x6B, 0x72, 0x80); DARK = RGBColor(0x23, 0x30, 0x3F); TEAL = RGBColor(0x1F, 0x6F, 0x6B)

    if limit:
        rows = conn.execute("SELECT * FROM items WHERE kind=? ORDER BY length(description) DESC LIMIT ?",
                            (kind, int(limit))).fetchall()
    else:
        rows = _rows(conn, kind)

    group_field = "vertical" if kind == "task" else "source"
    groups = {}
    for r in rows:
        groups.setdefault(r[group_field] or "Other", []).append(r)

    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"; normal.size = Pt(11)

    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Adobe Freelance Tasks" if kind == "task" else "Adobe Job Listings")
    tr.bold = True; tr.font.size = Pt(24); tr.font.color.rgb = INK
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("%d entries across %d %s" % (len(rows), len(groups), group_field + "s"))
    sr.font.size = Pt(12); sr.font.color.rgb = ACCENT
    doc.add_paragraph()

    accent_for = ACCENT if kind == "task" else TEAL
    for gname in sorted(groups, key=lambda g: -len(groups[g])):
        h = doc.add_heading(level=1)
        hr = h.add_run("%s  (%d)" % (gname, len(groups[gname]))); hr.font.color.rgb = INK
        for i, r in enumerate(groups[gname], 1):
            tools = "  ·  ".join(json.loads(r["adobe_tools"] or "[]")) or "Adobe Creative Cloud"
            p = doc.add_paragraph(); pr = p.add_run("%d.  %s%s" % (i, r["title"] or "", ("  —  " + r["company"]) if r["company"] else ""))
            pr.bold = True; pr.font.size = Pt(12.5); pr.font.color.rgb = INK
            meta = " · ".join([x for x in [r["source"], r["budget"] or r["salary"], r["location"], r["posted_at"]] if x])
            mp = doc.add_paragraph(); mr = mp.add_run(meta); mr.italic = True; mr.font.size = Pt(9.5); mr.font.color.rgb = GREY
            tp = doc.add_paragraph(); tl = tp.add_run("Adobe skills:  " + tools); tl.bold = True; tl.font.size = Pt(10.5); tl.font.color.rgb = accent_for
            if r["description"]:
                dp = doc.add_paragraph(); dr = dp.add_run(r["description"][:2400]); dr.font.size = Pt(10.5); dr.font.color.rgb = DARK
            lp = doc.add_paragraph()
            if r["url"]:
                _add_hyperlink(lp, r["url"], "View source →")
            doc.add_paragraph()

    doc.save(path)
    return len(rows)
